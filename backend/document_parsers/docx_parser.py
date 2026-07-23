from io import BytesIO

from docx import Document

from backend.document_parsers.parser_exceptions import InvalidDocumentError


def parse_docx(file_bytes: bytes) -> str:
    """
    Extract paragraphs and table contents from a DOCX document.
    """
    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise InvalidDocumentError(
            "The uploaded DOCX document could not be opened."
        ) from exc

    sections: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    for table in document.tables:
        table_rows: list[str] = []

        for row in table.rows:
            cell_values = [
                cell.text.strip()
                for cell in row.cells
            ]

            if any(cell_values):
                table_rows.append(" | ".join(cell_values))

        if table_rows:
            sections.append("\n".join(table_rows))

    return "\n\n".join(sections).strip()